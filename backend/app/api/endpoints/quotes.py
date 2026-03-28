"""Quote API endpoints."""

from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from typing import List, Optional
import uuid

from app.api.deps import get_db, get_current_active_user, require_role
from app.schemas.quote import (
    QuoteResponse, QuoteCreateRequest, QuoteAcceptResponse,
    QuoteForCustomerResponse, QuoteProviderInfo, QuoteDocExtractResponse
)
from app.models.user import User
from app.services.rfq_service import submit_quote, accept_quote

router = APIRouter()


@router.post("/provider/rfqs/{rfq_id}/quote/extract-document", response_model=QuoteDocExtractResponse)
async def extract_quote_document(
    rfq_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["provider"])),
):
    """Upload a document and use LLM to extract quote fields for pre-filling."""
    import io
    import json
    import mimetypes
    import logging
    from decimal import Decimal
    from app.services.config_service import get_runtime_config
    from app.services.file_service import generate_unique_key
    from app.core.config import settings
    import boto3
    from botocore.config import Config as BotoConfig

    logger = logging.getLogger(__name__)

    filename = file.filename or "upload"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in ("pdf", "docx", "txt"):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and TXT files are supported")

    file_bytes = await file.read()
    if len(file_bytes) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Maximum 10MB allowed.")

    # Extract text from document
    try:
        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            doc_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text: parts.append(cell.text)
            doc_text = "\n".join(parts)
        else:
            doc_text = file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to extract text: {e}")

    # Upload to S3
    s3_key = generate_unique_key("quote-documents", filename)
    try:
        s3 = boto3.client(
            "s3",
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_REGION,
            config=BotoConfig(signature_version="s3v4"),
        )
        mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
        s3.put_object(Bucket=settings.S3_BUCKET_NAME, Key=s3_key, Body=file_bytes, ContentType=mime_type)
    except Exception as e:
        logger.warning(f"S3 upload failed (continuing): {e}")
        s3_key = ""

    # LLM extraction
    cfg = await get_runtime_config(db)
    # Use LLM3 (Document Collapse LLM) with fallback to LLM2 (Firm Ranking LLM)
    llm_api_key = cfg.get("DOC_LLM_API_KEY") or cfg.get("OPENAI_API_KEY", "")
    llm_api_base = cfg.get("DOC_LLM_API_BASE") or cfg.get("OPENAI_API_BASE", "https://api.openai.com/v1")
    llm_model = cfg.get("DOC_LLM_MODEL") or cfg.get("OPENAI_LLM_MODEL", "gpt-4o-mini")

    extracted = {}
    raw_extraction = ""

    if llm_api_key and doc_text.strip():
        try:
            from openai import AsyncOpenAI
            client = AsyncOpenAI(api_key=llm_api_key, base_url=llm_api_base)
            system_prompt = (
                "You are an expert engineering quote analyst specializing in engineering services procurement. "
                "Parse the provided quote document and extract structured information. "
                "Return ONLY valid JSON with these exact keys (use null for missing values, never omit keys). "
                "Field extraction rules: "
                "rough_price_min: Minimum or lower-bound price as a plain number (no symbols). "
                "If document has a single total price, use that value for both min and max. "
                "Look for: Total, Subtotal, Grand Total, Estimate, Budget, Price, Cost. "
                "rough_price_max: Maximum or upper-bound price. Same source as min if single price. "
                "currency: ISO 4217 3-letter code (e.g. USD, EUR, CAD). Default USD if not stated. "
                "turnaround_estimate_text: Max 200 chars. Look for Lead Time, Delivery Schedule, "
                "Completion Date, Timeline, Weeks/Months from PO, ETA. Quote the text directly. "
                "assumptions_text: Max 1000 chars, formatted as bullet points starting with dash space. "
                "FIRST look for an explicit Assumptions or Basis of Design section. "
                "If none exists, INFER implicit assumptions from: technical standards referenced "
                "(ASME, AISC, ISA, API, IEC, ANSI, AWS, NEMA, IEEE, etc.), material or component specifications, "
                "environmental or site conditions implied, scope exclusions or boundaries stated, "
                "payment or schedule conditions, any clauses beginning with subject to, "
                "based on, assumes, provided that, or per customer. "
                "Always populate this field if engineering or technical context exists in the document. "
                "scope_notes: Max 1000 chars. Look for Scope of Work, Deliverables, Line Items, "
                "Services Included, Description of Work. List key deliverables and any explicit exclusions. "
                "Return only valid JSON. No markdown, no explanation."
            )
            response = await client.chat.completions.create(
                model=llm_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Extract quote fields from this document:\n\n{doc_text[:8000]}"},
                ],
                temperature=0,
                max_tokens=1000,
            )
            raw_extraction = response.choices[0].message.content or "{}"
            clean = raw_extraction.strip()
            if clean.startswith("```"):
                clean = clean.split("```")[1]
                if clean.startswith("json"): clean = clean[4:]
            extracted = json.loads(clean)
        except json.JSONDecodeError:
            extracted = {}
        except Exception as e:
            logger.error(f"LLM extraction failed: {e}")
            extracted = {}

    def safe_decimal(val):
        try: return Decimal(str(val)) if val is not None else None
        except Exception: return None

    return QuoteDocExtractResponse(
        s3_key=s3_key,
        original_filename=filename,
        extracted_fields=extracted,
        rough_price_min=safe_decimal(extracted.get("rough_price_min")),
        rough_price_max=safe_decimal(extracted.get("rough_price_max")),
        currency=str(extracted.get("currency", "USD") or "USD")[:3],
        turnaround_estimate_text=extracted.get("turnaround_estimate_text"),
        assumptions_text=extracted.get("assumptions_text"),
        scope_notes=extracted.get("scope_notes"),
        raw_extraction=raw_extraction,
    )


def _get_provider_email(provider) -> str:
    """Get first valid email from provider email_addresses field."""
    if not provider:
        return ""
    emails = provider.email_addresses
    if not emails:
        return ""
    import json
    if isinstance(emails, list):
        return emails[0] if emails else ""
    if isinstance(emails, str):
        try:
            parsed = json.loads(emails)
            if isinstance(parsed, list): return parsed[0] if parsed else ""
            return str(parsed)
        except Exception:
            return emails
    return ""


@router.post("/provider/rfqs/{rfq_id}/quote", response_model=QuoteResponse)
async def submit_provider_quote(
    rfq_id: str,
    data: QuoteCreateRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["provider"])),
):
    """Submit a quote for an unlocked RFQ (provider only)."""
    from sqlalchemy import select, func
    from app.models.provider import ProviderMembership
    from app.models.rfq import RFQ
    from app.models.quote import Quote
    from app.models.enums import RfqStatus
    from app.core.config import settings

    result = await db.execute(select(ProviderMembership).where(ProviderMembership.user_id == current_user.id))
    membership = result.scalar_one_or_none()
    if not membership:
        raise HTTPException(status_code=403, detail="No provider firm linked to your account")

    try:
        rfq_uuid = uuid.UUID(rfq_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid RFQ ID format")

    try:
        quote = await submit_quote(db=db, data=data, rfq_id=rfq_uuid, provider_id=membership.provider_id, user=current_user)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # Save optional document reference
    if data.document_s3_key:
        quote.document_s3_key = data.document_s3_key
        quote.document_filename = data.document_filename
        await db.commit()
        await db.refresh(quote)

    # Update quote_count using live SQL count
    rfq_result = await db.execute(select(RFQ).where(RFQ.id == rfq_uuid))
    rfq = rfq_result.scalar_one_or_none()
    if rfq:
        _live_res = await db.execute(
            select(func.count()).select_from(Quote).where(
                Quote.rfq_id == rfq_uuid,
                Quote.quote_status.in_(["submitted", "accepted"]),
            )
        )
        rfq.quote_count = _live_res.scalar() or 0
        max_quotes = getattr(settings, "RFQ_MAX_QUOTES", 5)
        if rfq.quote_count >= max_quotes:
            rfq.rfq_status = RfqStatus.QUOTE_LIMIT_REACHED
            rfq.is_closed = True
        await db.commit()
        await db.refresh(rfq)

    async def _notify_customer():
        try:
            from app.services.email_service import send_quote_notification
            from sqlalchemy.orm import joinedload
            from app.db.session import AsyncSessionLocal
            from sqlalchemy import select as _sel
            async with AsyncSessionLocal() as notify_db:
                q_result = await notify_db.execute(
                    _sel(Quote).options(joinedload(Quote.rfq), joinedload(Quote.provider)).where(Quote.id == quote.id)
                )
                full_quote = q_result.scalar_one_or_none()
                if full_quote and full_quote.rfq and full_quote.rfq.customer_email:
                    await send_quote_notification(db=notify_db, recipient_email=full_quote.rfq.customer_email, quote=full_quote)
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Quote notification email failed: {e}")

    background_tasks.add_task(_notify_customer)
    return QuoteResponse.from_orm(quote)


@router.get("/customer/rfqs/{rfq_id}/quotes", response_model=List[QuoteForCustomerResponse])
async def get_customer_quotes(
    rfq_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Get all quotes for customer RFQ with provider info. Accepted quotes include contact + document URL."""
    from sqlalchemy import select
    from sqlalchemy.orm import selectinload
    from app.models.rfq import RFQ
    from app.models.quote import Quote
    from app.services.file_service import generate_download_url

    try:
        rfq_uuid = uuid.UUID(rfq_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid RFQ ID")

    result = await db.execute(select(RFQ).where(RFQ.id == rfq_uuid))
    rfq = result.scalar_one_or_none()
    if not rfq:
        raise HTTPException(status_code=404, detail="RFQ not found")
    if rfq.customer_user_id != current_user.id and "admin" not in (current_user.roles or []):
        raise HTTPException(status_code=403, detail="Not authorized")

    result = await db.execute(
        select(Quote)
        .options(selectinload(Quote.files), selectinload(Quote.provider))
        .where(Quote.rfq_id == rfq_uuid)
        .order_by(Quote.created_at.desc())
    )
    quotes = result.scalars().all()

    responses = []
    for q in quotes:
        provider = q.provider
        is_accepted = q.quote_status == "accepted"

        provider_info = QuoteProviderInfo(
            provider_id=provider.id if provider else 0,
            provider_name=provider.name or provider.firm_name or "Unknown" if provider else "Unknown",
            firm_name=provider.firm_name or provider.name or "" if provider else "",
            primary_specialty=provider.primary_specialty if provider else None,
            website=provider.website if (provider and is_accepted) else None,
            email=_get_provider_email(provider) if (provider and is_accepted) else None,
            phone=provider.phone if (provider and is_accepted) else None,
            city=provider.city if (provider and is_accepted) else None,
            state=provider.state if (provider and is_accepted) else None,
            address=provider.address if (provider and is_accepted) else None,
        )

        # Generate download URL for document if accepted and has document
        doc_download_url = None
        if is_accepted and q.document_s3_key:
            try:
                doc_download_url = generate_download_url(q.document_s3_key, expire_seconds=3600)
            except Exception:
                pass

        quote_dict = {
            "id": q.id,
            "rfq_id": q.rfq_id,
            "created_at": q.created_at,
            "updated_at": q.updated_at,
            "provider_id": q.provider_id,
            "submitter_user_id": q.submitter_user_id,
            "quote_status": q.quote_status,
            "rough_price_min": q.rough_price_min,
            "rough_price_max": q.rough_price_max,
            "currency": q.currency,
            "turnaround_estimate_text": q.turnaround_estimate_text,
            "assumptions_text": q.assumptions_text,
            "scope_notes": q.scope_notes,
            "submitted_at": q.submitted_at,
            "customer_viewed_at": q.customer_viewed_at,
            "document_s3_key": q.document_s3_key if is_accepted else None,
            "document_filename": q.document_filename if is_accepted else None,
            "customer_contact_name": None,
            "customer_company": None,
            "customer_email": None,
            "provider": provider_info,
            "files": [{
                "id": f.id, "quote_id": f.quote_id,
                "original_filename": f.original_filename,
                "mime_type": f.mime_type,
                "file_size_bytes": f.file_size_bytes,
                "created_at": f.created_at,
            } for f in (q.files or [])],
            "document_download_url": doc_download_url,
        }
        responses.append(QuoteForCustomerResponse(**quote_dict))

    return responses



@router.get("/customer/quotes/{quote_id}/document")
async def get_quote_document_download(
    quote_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    from sqlalchemy import select
    from app.models.quote import Quote
    from app.models.rfq import RFQ
    from app.services.file_service import generate_download_url
    try:
        quote_uuid = uuid.UUID(quote_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid quote ID")
    result = await db.execute(select(Quote).where(Quote.id == quote_uuid))
    quote = result.scalar_one_or_none()
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    rfq_result = await db.execute(select(RFQ).where(RFQ.id == quote.rfq_id))
    rfq = rfq_result.scalar_one_or_none()
    if not rfq:
        raise HTTPException(status_code=404, detail="RFQ not found")
    if rfq.customer_user_id != current_user.id and "admin" not in (current_user.roles or []):
        raise HTTPException(status_code=403, detail="Not authorized")
    if quote.quote_status != "accepted":
        raise HTTPException(status_code=403, detail="Document only available for accepted quotes")
    if not quote.document_s3_key:
        raise HTTPException(status_code=404, detail="No document attached to this quote")
    try:
        url = generate_download_url(quote.document_s3_key, expire_seconds=3600)
        return {"download_url": url, "filename": quote.document_filename or "quote-document"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate download URL: {e}")


@router.post("/customer/quotes/{quote_id}/accept", response_model=QuoteAcceptResponse)
async def accept_quote_endpoint(
    quote_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    try:
        quote_uuid = uuid.UUID(quote_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid quote ID")
    try:
        provider_contact = await accept_quote(db, quote_uuid, current_user)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e))
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Accept quote error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    from sqlalchemy import select as _select
    from app.models.quote import Quote as _Quote
    result = await db.execute(_select(_Quote).where(_Quote.id == quote_uuid))
    quote = result.scalar_one_or_none()
    # Determine message based on NDA status
    nda_triggered = provider_contact.get("nda_triggered")
    nda_error = provider_contact.get("nda_error")
    if nda_triggered is True:
        msg = "Quote accepted. Both parties will receive an NDA to sign via email."
    elif nda_triggered is False and nda_error:
        msg = f"Quote accepted, but NDA sending failed: {nda_error}. Please contact support."
    else:
        msg = "Quote accepted. Provider contact details are now revealed."

    return QuoteAcceptResponse(
        success=True,
        message=msg,
        rfq_id=quote.rfq_id,
        selected_quote_id=quote.id,
        selected_provider_id=quote.provider_id,
        provider_contact_revealed=True,
        provider_name=provider_contact.get("provider_name"),
        provider_email=provider_contact.get("provider_email"),
        provider_phone=provider_contact.get("provider_phone"),
        provider_website=provider_contact.get("provider_website"),
        provider_city=provider_contact.get("provider_city"),
        provider_state=provider_contact.get("provider_state"),
        provider_address=provider_contact.get("provider_address"),
        nda_triggered=nda_triggered,
        nda_error=nda_error,
    )


@router.post("/provider/quotes/{quote_id}/withdraw")
async def withdraw_quote(
    quote_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["provider"])),
):
    from sqlalchemy import select
    from app.models.quote import Quote
    from app.models.provider import ProviderMembership
    result = await db.execute(select(Quote).where(Quote.id == quote_id))
    quote = result.scalar_one_or_none()
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    result = await db.execute(
        select(ProviderMembership).where(
            ProviderMembership.provider_id == quote.provider_id,
            ProviderMembership.user_id == current_user.id,
        )
    )
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=403, detail="Not authorized")
    quote.quote_status = "withdrawn"
    await db.commit()
    return {"message": "Quote withdrawn"}


@router.get("/provider/quotes/me", response_model=List[QuoteResponse])
async def get_provider_quotes(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["provider"])),
):
    from sqlalchemy import select
    from app.models.quote import Quote
    from app.models.provider import ProviderMembership
    from app.models.rfq import RFQ
    result = await db.execute(select(ProviderMembership).where(ProviderMembership.user_id == current_user.id))
    membership = result.scalar_one_or_none()
    if not membership:
        return []
    result = await db.execute(select(Quote).where(Quote.provider_id == membership.provider_id))
    quotes = result.scalars().all()
    responses = []
    for q in quotes:
        resp = QuoteResponse.from_orm(q)
        if q.quote_status == "accepted":
            rfq_result = await db.execute(select(RFQ).where(RFQ.id == q.rfq_id))
            rfq = rfq_result.scalar_one_or_none()
            if rfq:
                resp.customer_contact_name = rfq.contact_name
                resp.customer_company = rfq.business_name
                resp.customer_email = rfq.customer_email
        responses.append(resp)
    return responses
