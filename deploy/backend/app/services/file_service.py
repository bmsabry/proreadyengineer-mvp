"""File storage service with S3 presigned URLs and document text extraction."""

import io
import mimetypes
import uuid
from datetime import timedelta
from typing import Any, Optional
from urllib.parse import urlparse

import boto3
from botocore.config import Config
from botocore.exceptions import ClientError

from app.core.config import settings


def _get_s3_client() -> boto3.client:
    """Get configured S3 client.

    Returns:
        boto3.client: S3 client instance.
    """
    return boto3.client(
        "s3",
        aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
        aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
        region_name=settings.AWS_REGION,
        config=Config(signature_version="s3v4"),
    )


def generate_upload_url(
    key: str,
    content_type: str,
    expire_seconds: int = 3600,
    max_file_size: int = 26214400,  # 25MB default
) -> dict[str, Any]:
    """Generate S3 presigned POST URL for file upload.

    Args:
        key: S3 object key (path in bucket).
        content_type: MIME type of file.
        expire_seconds: URL expiration time in seconds.
        max_file_size: Maximum file size in bytes.

    Returns:
        dict: Contains 'url', 'fields' for POST, and 'key'.

    Raises:
        RuntimeError: If URL generation fails.
    """
    s3 = _get_s3_client()

    try:
        # Generate presigned POST
        conditions = [
            ["content-length-range", 0, max_file_size],
            ["eq", "$Content-Type", content_type],
        ]

        response = s3.generate_presigned_post(
            Bucket=settings.S3_BUCKET_NAME,
            Key=key,
            Fields={
                "Content-Type": content_type,
            },
            Conditions=conditions,
            ExpiresIn=expire_seconds,
        )

        return {
            "url": response["url"],
            "fields": response["fields"],
            "key": key,
            "expires_in": expire_seconds,
        }
    except ClientError as e:
        raise RuntimeError(f"Failed to generate upload URL: {e}")


def generate_download_url(
    key: str,
    expire_seconds: int = 3600,
    inline: bool = False,
) -> str:
    """Generate S3 presigned GET URL for file download.

    Args:
        key: S3 object key.
        expire_seconds: URL expiration time in seconds.
        inline: Whether to display inline (True) or download (False).

    Returns:
        str: Presigned URL.

    Raises:
        RuntimeError: If URL generation fails.
    """
    s3 = _get_s3_client()

    try:
        # Determine content disposition
        content_disposition = "inline" if inline else "attachment"

        params = {
            "Bucket": settings.S3_BUCKET_NAME,
            "Key": key,
        }

        # Add response content disposition
        if not inline:
            filename = key.split("/")[-1]
            params["ResponseContentDisposition"] = f'attachment; filename="{filename}"'

        url = s3.generate_presigned_url(
            "get_object",
            Params=params,
            ExpiresIn=expire_seconds,
        )

        return url
    except ClientError as e:
        raise RuntimeError(f"Failed to generate download URL: {e}")


def delete_file(key: str) -> bool:
    """Delete a file from S3.

    Args:
        key: S3 object key.

    Returns:
        bool: True if deleted successfully.
    """
    s3 = _get_s3_client()

    try:
        s3.delete_object(
            Bucket=settings.S3_BUCKET_NAME,
            Key=key,
        )
        return True
    except ClientError:
        return False


def check_file_exists(key: str) -> bool:
    """Check if a file exists in S3.

    Args:
        key: S3 object key.

    Returns:
        bool: True if file exists.
    """
    s3 = _get_s3_client()

    try:
        s3.head_object(
            Bucket=settings.S3_BUCKET_NAME,
            Key=key,
        )
        return True
    except ClientError as e:
        if e.response["Error"]["Code"] == "404":
            return False
        raise


def generate_unique_key(
    prefix: str,
    filename: str,
    entity_id: Optional[uuid.UUID] = None,
) -> str:
    """Generate unique S3 key with organized path structure.

    Args:
        prefix: Path prefix (e.g., 'rfqs', 'quotes', 'ads').
        filename: Original filename.
        entity_id: Optional entity UUID for organization.

    Returns:
        str: Unique S3 key.
    """
    # Generate unique ID
    unique_id = str(uuid.uuid4())

    # Sanitize filename
    safe_filename = "".join(
        c for c in filename if c.isalnum() or c in (".", "-", "_")
    ).rstrip()

    # Build path
    if entity_id:
        key = f"{prefix}/{entity_id}/{unique_id}/{safe_filename}"
    else:
        key = f"{prefix}/{unique_id}/{safe_filename}"

    return key


async def extract_document_text(s3_key: str) -> str:
    """Extract text content from a document stored in S3.

    Supports PDF and DOCX files.

    Args:
        s3_key: S3 object key.

    Returns:
        str: Extracted text content.

    Raises:
        ValueError: If file type not supported.
        RuntimeError: If extraction fails.
    """
    # Determine file type
    ext = s3_key.lower().split(".")[-1] if "." in s3_key else ""

    # Download file from S3
    s3 = _get_s3_client()

    try:
        response = s3.get_object(
            Bucket=settings.S3_BUCKET_NAME,
            Key=s3_key,
        )
        file_content = response["Body"].read()
    except ClientError as e:
        raise RuntimeError(f"Failed to download file from S3: {e}")

    # Extract based on file type
    if ext == "pdf":
        return await _extract_pdf_text(file_content)
    elif ext == "docx":
        return await _extract_docx_text(file_content)
    elif ext in ["txt", "md", "csv"]:
        return file_content.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {ext}")


async def _extract_pdf_text(file_content: bytes) -> str:
    """Extract text from PDF content.

    Args:
        file_content: PDF file bytes.

    Returns:
        str: Extracted text.
    """
    try:
        import PyPDF2

        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text_parts = []
        for page in pdf_reader.pages:
            text_parts.append(page.extract_text() or "")

        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF text: {e}")


async def _extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX content.

    Args:
        file_content: DOCX file bytes.

    Returns:
        str: Extracted text.
    """
    try:
        from docx import Document

        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)

        text_parts = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)

        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX text: {e}")


def get_mime_type(filename: str) -> str:
    """Get MIME type from filename.

    Args:
        filename: Name of file.

    Returns:
        str: MIME type or 'application/octet-stream'.
    """
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or "application/octet-stream"


def validate_file_type(
    filename: str,
    allowed_extensions: Optional[list[str]] = None,
) -> bool:
    """Validate file extension against allowed types.

    Args:
        filename: Name of file.
        allowed_extensions: List of allowed extensions (e.g., ['.pdf', '.docx']).

    Returns:
        bool: True if valid.
    """
    if not allowed_extensions:
        allowed_extensions = settings.ALLOWED_UPLOAD_EXTENSIONS

    ext = filename.lower().split(".")[-1] if "." in filename else ""
    return f".{ext}" in allowed_extensions
